from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "public"
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(.45); sec.bottom_margin = Inches(.45)
sec.left_margin = Inches(.58); sec.right_margin = Inches(.58)
styles = doc.styles
styles["Normal"].font.name = "Arial"; styles["Normal"].font.size = Pt(8.6)
styles["Normal"].font.color.rgb = RGBColor.from_string("202020")
styles["Normal"].paragraph_format.space_after = Pt(2)

def p(text="", size=8.6, bold=False, color="202020", before=0, after=2, align=None):
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(before); par.paragraph_format.space_after = Pt(after)
    if align: par.alignment = align
    run = par.add_run(text); run.bold = bold; run.font.name = "Arial"; run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color)
    return par

def heading(text):
    par = p(text.upper(), 9.5, True, "111111", 8, 3)
    par.paragraph_format.keep_with_next = True
    par.paragraph_format.bottom_border = None
    return par

def job(title, dates, bullets):
    table = doc.add_table(rows=1, cols=2); table.autofit = False
    table.columns[0].width = Inches(5.8); table.columns[1].width = Inches(1.2)
    a, b = table.rows[0].cells
    a.text = title; b.text = dates
    for cell in (a,b):
        cell.vertical_alignment = 1
        for r in cell.paragraphs[0].runs: r.bold=True; r.font.name="Arial"; r.font.size=Pt(8.8)
    b.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for item in bullets:
        par = doc.add_paragraph(style="List Bullet")
        par.paragraph_format.left_indent = Inches(.16); par.paragraph_format.first_line_indent = Inches(-.12); par.paragraph_format.space_after = Pt(1)
        run = par.add_run(item); run.font.name="Arial"; run.font.size=Pt(8.1)

p("LUCAS FRISCHEISEN", 23, True, "111111", 0, 0)
p("ENGENHEIRO DE IA  •  BACKEND  •  SÓCIO-FUNDADOR", 9.5, True, "E5365B", 0, 5)
p("Jundiaí/SP  |  +55 11 98182-6659  |  lucas.frischeisen@gmail.com", 8.3, False, "444444", 0, 1)
p("linkedin.com/in/rukafuu  |  github.com/Rukafuu  |  waifucorp.org", 8.3, False, "444444", 0, 4)

heading("Resumo")
p("Engenheiro de IA, backend e sócio-fundador da wAIfu Corp, com experiência em IA generativa, agentes autônomos, RAG, voz e arquitetura de produtos digitais. Atua da pesquisa à produção, combinando TypeScript/Node.js, Python, Rust e infraestrutura cloud com foco em sistemas confiáveis e expressivos.", 8.5, False, "303030", 0, 2)

heading("Experiência")
job("Sócio-fundador & Engenheiro de IA  |  wAIfu Corp", "2026 - Atual", [
    "Lucas (Rukafuusca): cofundador de um coletivo que transforma experimentos de IA em software open source, incluindo Above All Graphs, Bastion e DRAGON.",
])
job("Engenheiro de IA  |  Heon", "08/06/2026 - Atual", [
    "Lidera a arquitetura e a evolução da Lia, IA multifuncional para clínicas integrada à plataforma Heon.",
    "Desenvolve serviços backend em TypeScript/Node.js e orquestrações n8n, com integrações via WhatsApp, Discord e APIs PHP.",
    "Projeta fluxos multiagentes, estratégias de fine-tuning/RLAIF e modos de atuação ativo e passivo.",
    "Contribui para a arquitetura do BFF Fastify do Heon Aluno, incluindo autenticação, permissões e contratos Zod para um público estimado de 6 mil usuários.",
])
job("Engenheiro de IA & Full Stack  |  Projetos independentes", "01/2024 - Atual", [
    "Criação da LiraVtuber, assistente multimodal com voz, memória, visão, Live2D e integrações em tempo real.",
    "Pesquisa e desenvolvimento do CAFUNÉ, arquitetura híbrida SNN 11D + Transformer bidirecional com aproximadamente 22,5M de parâmetros.",
    "Construção de Raegis para auditoria de modelos, fidelidade de RAG, entropia e detecção de anomalias.",
])
job("T&D / Automação de Processos  |  Xiaomi", "06/2024 - 02/2026", [
    "Automação de processos, desenvolvimento de materiais técnicos e capacitação de equipes com foco em eficiência operacional.",
])
job("Supervisor de Vendas  |  Samsung", "01/2022 - 04/2024", [
    "Liderança de equipe, acompanhamento de indicadores e desenvolvimento de pessoas em ambiente de alta performance.",
])

heading("Formação")
p("Engenharia de Inteligência Artificial — FMU  |  conclusão prevista: 06/2028", 8.5, True, "202020", 0, 1)
p("Desenvolvimento de Sistemas — ETEC  |  concluído: 2025", 8.5, True, "202020", 0, 2)

heading("Competências")
p("IA & Dados: PyTorch, Transformers, LangChain, Gemini, TensorFlow, Scikit-learn, RAG, fine-tuning, STT/TTS, RLAIF", 8.15)
p("Engenharia: TypeScript, JavaScript, Node.js, Fastify, Python, FastAPI, Django, React, Rust, Docker, CI/CD, Linux", 8.15)
p("Dados & Cloud: SQL, NoSQL, Redis, Firebase, GCP, Vercel  |  Idiomas: Inglês C1, Espanhol B2, Japonês A2", 8.15)

OUT.mkdir(exist_ok=True)
docx = OUT / "curriculo-lucas-frischeisen.docx"
doc.save(docx)
print(docx)
